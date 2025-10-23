"""
Word Document Template Generator.

This module creates structured Word documents for financial reports:
1. Executive Summary section
2. Financial Metrics tables
3. Analysis and Commentary sections
4. Charts and visualizations (placeholders)
"""

from typing import List, Dict, Optional
from datetime import date
from decimal import Decimal
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from loguru import logger

from src.models.schemas import FinancialMetric


class WordTemplateGenerator:
    """
    Generate Word documents from financial data.
    
    Creates professional financial reports with:
    - Cover page
    - Executive summary
    - Financial metrics tables
    - Analysis sections
    - Appendices
    """
    
    def __init__(self):
        self.logger = logger.bind(component="word_template_generator")
        self.doc = None
    
    def create_financial_report(
        self,
        company_name: str,
        report_period: date,
        metrics: List[FinancialMetric],
        commentary: Dict[str, str],
        output_path: str
    ) -> str:
        """
        Create complete financial report.
        
        Args:
            company_name: Company name
            report_period: Reporting period end date
            metrics: List of financial metrics
            commentary: Dictionary of commentary sections
            output_path: Output file path
        
        Returns:
            Path to generated document
        """
        self.logger.info(
            "creating_financial_report",
            company=company_name,
            period=str(report_period)
        )
        
        # Create new document
        self.doc = Document()
        self._setup_styles()
        
        # Add sections
        self._add_cover_page(company_name, report_period)
        self._add_page_break()
        
        self._add_executive_summary(commentary.get("executive_summary", ""))
        self._add_page_break()
        
        self._add_financial_metrics_section(metrics)
        self._add_page_break()
        
        self._add_analysis_sections(commentary)
        
        # Save document
        self.doc.save(output_path)
        
        self.logger.info("financial_report_created", path=output_path)
        
        return output_path
    
    def _setup_styles(self):
        """Setup document styles."""
        styles = self.doc.styles
        
        # Title style
        if 'Custom Title' not in styles:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading 1 style modification
        heading1 = styles['Heading 1']
        heading1.font.name = 'Arial'
        heading1.font.size = Pt(16)
        heading1.font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading 2 style modification
        heading2 = styles['Heading 2']
        heading2.font.name = 'Arial'
        heading2.font.size = Pt(14)
        heading2.font.color.rgb = RGBColor(0, 102, 204)
    
    def _add_cover_page(self, company_name: str, report_period: date):
        """Add cover page."""
        # Company name
        title = self.doc.add_paragraph(company_name, style='Custom Title')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # Spacing
        
        # Report title
        subtitle = self.doc.add_paragraph('Financial Analysis Report', style='Heading 1')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # Spacing
        
        # Period
        period_text = f"Period Ending: {report_period.strftime('%B %d, %Y')}"
        period_para = self.doc.add_paragraph(period_text)
        period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period_para.runs[0].font.size = Pt(12)
        
        self.doc.add_paragraph()  # Spacing
        
        # Generation date
        from datetime import datetime
        gen_date = f"Generated: {datetime.now().strftime('%B %d, %Y')}"
        date_para = self.doc.add_paragraph(gen_date)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size = Pt(10)
        date_para.runs[0].font.italic = True
    
    def _add_executive_summary(self, summary_text: str):
        """Add executive summary section."""
        self.doc.add_heading('Executive Summary', level=1)
        
        if summary_text:
            # Split into paragraphs
            paragraphs = summary_text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = self.doc.add_paragraph(para_text.strip())
                    para.runs[0].font.size = Pt(11)
        else:
            self.doc.add_paragraph("No executive summary available.")
    
    def _add_financial_metrics_section(self, metrics: List[FinancialMetric]):
        """Add financial metrics tables."""
        self.doc.add_heading('Financial Metrics', level=1)
        
        # Group metrics by category
        revenue_metrics = [m for m in metrics if 'revenue' in m.metric_name.lower()]
        profit_metrics = [m for m in metrics if 'profit' in m.metric_name.lower() or 'income' in m.metric_name.lower()]
        balance_metrics = [m for m in metrics if 'asset' in m.metric_name.lower() or 'liability' in m.metric_name.lower() or 'equity' in m.metric_name.lower()]
        
        # Revenue metrics table
        if revenue_metrics:
            self.doc.add_heading('Revenue Metrics', level=2)
            self._add_metrics_table(revenue_metrics)
            self.doc.add_paragraph()  # Spacing
        
        # Profitability metrics table
        if profit_metrics:
            self.doc.add_heading('Profitability Metrics', level=2)
            self._add_metrics_table(profit_metrics)
            self.doc.add_paragraph()  # Spacing
        
        # Balance sheet metrics table
        if balance_metrics:
            self.doc.add_heading('Balance Sheet Metrics', level=2)
            self._add_metrics_table(balance_metrics[:10])  # Limit to 10
            self.doc.add_paragraph()  # Spacing
    
    def _add_metrics_table(self, metrics: List[FinancialMetric]):
        """Add a table of metrics."""
        if not metrics:
            return
        
        # Group by period
        from collections import defaultdict
        by_period = defaultdict(list)
        for m in metrics:
            by_period[m.period_end_date].append(m)
        
        # Get sorted periods
        periods = sorted(by_period.keys(), reverse=True)
        
        # Create table
        table = self.doc.add_table(rows=1, cols=len(periods) + 1)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        for i, period in enumerate(periods):
            header_cells[i + 1].text = period.strftime('%Y-%m-%d')
        
        # Get unique metric names
        metric_names = list(set(m.metric_name for m in metrics))
        
        # Add data rows
        for metric_name in metric_names:
            row_cells = table.add_row().cells
            row_cells[0].text = metric_name
            
            for i, period in enumerate(periods):
                # Find metric for this period
                metric = next((m for m in by_period[period] if m.metric_name == metric_name), None)
                if metric:
                    value_text = f"{metric.value:,.2f}"
                    row_cells[i + 1].text = value_text
                else:
                    row_cells[i + 1].text = "N/A"
    
    def _add_analysis_sections(self, commentary: Dict[str, str]):
        """Add analysis commentary sections."""
        self.doc.add_heading('Financial Analysis', level=1)
        
        sections = [
            ("Revenue Analysis", "revenue_analysis"),
            ("Profitability Analysis", "profitability_analysis"),
            ("Balance Sheet Analysis", "balance_sheet_analysis"),
            ("Cash Flow Analysis", "cash_flow_analysis"),
        ]
        
        for title, key in sections:
            content = commentary.get(key, "")
            if content:
                self.doc.add_heading(title, level=2)
                
                # Split into paragraphs
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = self.doc.add_paragraph(para_text.strip())
                        para.runs[0].font.size = Pt(11)
                
                self.doc.add_paragraph()  # Spacing
    
    def _add_page_break(self):
        """Add page break."""
        self.doc.add_page_break()
    
    def add_custom_section(self, title: str, content: str):
        """
        Add a custom section to the document.
        
        Args:
            title: Section title
            content: Section content
        """
        if not self.doc:
            self.doc = Document()
        
        self.doc.add_heading(title, level=1)
        
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                self.doc.add_paragraph(para_text.strip())
    
    def add_table(self, headers: List[str], data: List[List[str]]):
        """
        Add a custom table to the document.
        
        Args:
            headers: Table header row
            data: Table data rows
        """
        if not self.doc:
            self.doc = Document()
        
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # Data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
